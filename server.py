from flask import Flask, request, render_template, jsonify
import os
from docx import Document
from PyPDF2 import PdfReader
import requests as req

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('alpine.html')

@app.route('/rate', methods=['POST'])
def rate() -> jsonify:
    try:
        f = request.files.get('file')
        name = request.form.get('name')  
        age = request.form.get('age')   
        
        if f is not None and f.filename != '':
            if not os.path.exists('file'):
                os.makedirs('file')
            file_path = f"./file/{f.filename}"
            f.save(file_path)
            print(f"File saved: {f.filename} at {file_path}")
            file_extension = f.filename.split('.')[-1].lower()
            match file_extension:
                case 'pdf': pdf_text(PdfReader(file_path))
                case 'docx': docx_text(Document(file_path))
                case 'txt': 
                    with open(file_path, 'r'):
                        text = f.read()
                case _: 
                    return jsonify({
                        "status": "error"
                    })
                
            return jsonify({
                "status": "success"
            })
            
        else:
            print("No file uploaded or empty filename")
            return jsonify({
                "status": "error"
            })
        
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

def docx_text(docx: Document) -> str:
    text: str = ""
    for paragraph in docx.paragraphs:
        text += paragraph.text
    return text

def pdf_text(pdf: PdfReader) -> str:
    text: str = ""
    page_content = pdf.pages
    for i in page_content:
        text += i.extract_text()
    return text


def response(query: str)-> str:
    payload = {
        "model": "NorthEye:1.0",
        "prompt": query,
        "stream": False
    }
    response = req.post("http://localhost:11434/api/generate", headers={"Content-Type": "application/json"}, json=payload)
    if response.status_code == 200:
        print('Success') #debug
        return response.json()
    else:
        print(f"Error: {response.status_code}")  #debug
        return None

@app.route('/')
def index() -> render_template:
    return render_template('alpine.html')

@app.route('/rate')
def rate():
    data = request.get_json()
    job_description = data.get('description', '')
    qualifications = data.get('qualifications', '')
    resume = data.get('resume', '')
    ext: str = file_check(resume)
    resume_text = ""

    for i in data: #debug
        print(i)   #debug

    match(ext.lower()):
        case 'pdf': resume_text += pdf_text(PdfReader(resume)) 
        case 'docx': resume_text += docx_text(Document(resume))
        case _: resume_text += ""

    prompt = f"""
            {{
                Job Descriptions: {job_description.strip().replace('\n', '  ')},
                Qualifications: {qualifications.strip().replace('\n', '  ')},
                Resume: {resume_text.strip().replace('\n', '  ')}
            }}
    """

    res = response(prompt)
    with open('res.json', 'w') as f:    #debug
        f.write(res)                    #debug

    return app.response_class(
        response=res,
        mimetype='application/json'
    )

if __name__ == '__main__':
    app.run(port=9000, host='localhost', debug=True)